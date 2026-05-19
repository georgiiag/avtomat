import os
import time
from typing import Optional

from flask import Flask, render_template, jsonify, request, Response, redirect, url_for, send_file
from werkzeug.utils import secure_filename

app = Flask(__name__)

def load_tooling_registry_seed():
    try:
        import openpyxl
    except Exception:
        return []

    fname = "Перечень_оборудования_оснастки_V4_6_8_Всё с оснасткой_ТО (1).xlsx"
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    xlsx_path = os.path.join(project_root, fname)
    if not os.path.exists(xlsx_path):
        return []

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception:
        return []

    def norm(v) -> str:
        return str(v or "").strip().lower().replace("\xa0", " ")

    def find_headers(ws):
        max_scan_rows = min(120, int(getattr(ws, "max_row", 0) or 0) or 0)
        max_scan_cols = min(80, int(getattr(ws, "max_column", 0) or 0) or 0)
        if not max_scan_rows or not max_scan_cols:
            return None

        for r in range(1, max_scan_rows + 1):
            op_col = None
            eq_col = None
            inv_col = None

            for c in range(1, max_scan_cols + 1):
                s = norm(ws.cell(r, c).value)
                if not s:
                    continue

                if inv_col is None and "инвентар" in s:
                    inv_col = c

                if eq_col is None and "оборуд" in s and ("наимен" in s or "тип" in s or "марк" in s or "модел" in s):
                    eq_col = c

                if op_col is None and (("операц" in s) or ("технолог" in s and ("операц" in s or "осуществ" in s))):
                    op_col = c

            if op_col and eq_col and inv_col:
                return (r, op_col, eq_col, inv_col)

        return None

    ws = None
    header_row = None
    op_col = None
    eq_col = None
    inv_col = None

    for sn in getattr(wb, "sheetnames", []) or []:
        candidate_ws = wb[sn]
        found = find_headers(candidate_ws)
        if found:
            ws = candidate_ws
            header_row, op_col, eq_col, inv_col = found
            break

    if ws is None:
        try:
            ws = wb.active
        except Exception:
            return []
        found = find_headers(ws)
        if not found:
            return []
        header_row, op_col, eq_col, inv_col = found

    items = []
    seen = set()
    empty_streak = 0
    max_rows = int(getattr(ws, "max_row", 0) or 0) or 0

    for r in range(header_row + 1, max_rows + 1):
        op_raw = ws.cell(r, op_col).value if op_col else ""
        eq_raw = ws.cell(r, eq_col).value if eq_col else ""
        inv_raw = ws.cell(r, inv_col).value if inv_col else ""
        op = str(op_raw).strip() if op_raw is not None else ""
        eq = str(eq_raw).strip() if eq_raw is not None else ""
        inv = str(inv_raw).strip() if inv_raw is not None else ""

        if not (op or eq or inv):
            empty_streak += 1
            if empty_streak >= 25:
                break
            continue
        empty_streak = 0

        if not (eq or inv):
            continue

        if not inv:
            inv = f"INV-{len(items) + 1:04d}"

        base = inv
        i = 1
        while inv in seen:
            i += 1
            inv = f"{base}-{i}"
        seen.add(inv)

        items.append({
            "id": inv,
            "name": eq,
            "description": op
        })

    return items

@app.after_request
def add_cors_headers(resp: Response):
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return resp

@app.route('/')
def index():
    return redirect(url_for("tp_constructor"))

@app.route('/index.html')
def index_html():
    return render_template('index.html')

@app.route('/tp_constructor')
@app.route('/tp_constructor.html')
def tp_constructor():
    return render_template('tp_constructor.html', tooling_registry=load_tooling_registry_seed())

@app.route('/api/tooling_seed_debug')
def tooling_seed_debug():
    data = load_tooling_registry_seed()
    return jsonify({"ok": True, "count": len(data), "sample": data[:3]})

@app.route('/recognizer')
@app.route('/recognizer.html')
def recognizer():
    return render_template('recognizer.html')

UPLOAD_DIR = os.path.join(app.root_path, "static", "shared_docs")
SNAPSHOT_DIR = os.path.join(app.root_path, "static", "shared_docs", "snapshots")
ATTACHMENTS_DIR = os.path.join(app.root_path, "static", "attachments")

def ensure_upload_dir():
    try:
        os.makedirs(UPLOAD_DIR, exist_ok=True)
    except Exception:
        pass

def ensure_snapshot_dir():
    try:
        os.makedirs(SNAPSHOT_DIR, exist_ok=True)
    except Exception:
        pass

def ensure_attachments_dir(subdir: Optional[str] = None):
    try:
        os.makedirs(ATTACHMENTS_DIR, exist_ok=True)
        if subdir:
            os.makedirs(os.path.join(ATTACHMENTS_DIR, subdir), exist_ok=True)
    except Exception:
        pass

def docx_to_simple_html(path: str) -> str:
    try:
        from docx import Document
    except Exception:
        return "<div style='color:#b13b2e'>python-docx не установлен</div>"
    try:
        doc = Document(path)
    except Exception as e:
        return f"<div style='color:#b13b2e'>Не удалось открыть DOCX: {e}</div>"

    parts = []
    for p in getattr(doc, "paragraphs", []):
        txt = (p.text or "").strip()
        if not txt:
            continue
        parts.append(f"<p>{txt}</p>")

    for t in getattr(doc, "tables", []):
        try:
            rows = []
            for r in t.rows:
                cells = []
                for c in r.cells:
                    cells.append(f"<td>{(c.text or '').strip()}</td>")
                rows.append("<tr>" + "".join(cells) + "</tr>")
            parts.append("<table>" + "".join(rows) + "</table>")
        except Exception:
            continue

    if not parts:
        return "<div style='color:#667085'>Документ пустой</div>"
    return "\n".join(parts)

@app.route('/share', methods=['GET', 'POST', 'OPTIONS'])
def share_doc():
    if request.method == "OPTIONS":
        return ("", 204)

    ensure_upload_dir()

    if request.method == "POST":
        if 'file' not in request.files:
            return jsonify({"ok": False, "error": "file is required"}), 400
        f = request.files['file']
        if not f or not f.filename:
            return jsonify({"ok": False, "error": "empty file"}), 400
        filename = secure_filename(f.filename)
        if not filename.lower().endswith(".docx"):
            return jsonify({"ok": False, "error": "only .docx supported"}), 400

        doc_id = f"doc-{int(__import__('time').time())}-{os.urandom(3).hex()}"
        save_name = f"{doc_id}__{filename}"
        path = os.path.join(UPLOAD_DIR, save_name)
        f.save(path)
        return redirect(url_for("view_shared_doc", doc_id=doc_id))

    files = []
    try:
        for name in os.listdir(UPLOAD_DIR):
            if "__" not in name or not name.lower().endswith(".docx"):
                continue
            doc_id, orig = name.split("__", 1)
            files.append((doc_id, orig, name))
        files.sort(key=lambda x: x[0], reverse=True)
    except Exception:
        files = []

    items = "\n".join([f"<li><a href='/share/{doc_id}'>{orig}</a></li>" for doc_id, orig, _ in files[:50]]) or "<li style='color:#667085'>Пока нет файлов</li>"

    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Общий доступ — DOCX</title>
  <style>
    body {{ font-family: Arial, sans-serif; background:#f0f2f5; margin:0; color:#111; }}
    .bar {{ background:#2c3e50; color:#fff; padding:14px 18px; display:flex; justify-content:space-between; }}
    .wrap {{ max-width:1000px; margin:0 auto; padding:18px; }}
    .card {{ background:#fff; border-radius:12px; padding:16px; box-shadow:0 4px 6px rgba(0,0,0,0.05); margin-bottom:14px; }}
    .btn {{ border:0; background:#3498db; color:#fff; padding:10px 14px; border-radius:10px; cursor:pointer; font-weight:600; }}
    input[type=file] {{ padding:8px; border:1px solid #dfe6e9; border-radius:10px; background:#fff; }}
    ul {{ margin:0; padding-left:18px; }}
    a {{ color:#2c3e50; }}
  </style>
</head>
<body>
  <div class="bar">
    <div>Общий доступ: загрузка DOCX для просмотра</div>
    <div><a href="/tp_constructor" style="color:#fff; text-decoration:none; opacity:0.9;">Конструктор ТП</a></div>
  </div>
  <div class="wrap">
    <div class="card">
      <form method="post" enctype="multipart/form-data">
        <div style="font-weight:700; margin-bottom:8px;">Загрузить DOCX</div>
        <div style="display:flex; gap:10px; align-items:center; flex-wrap:wrap;">
          <input name="file" type="file" accept=".docx,application/vnd.openxmlformats-officedocument.wordprocessingml.document">
          <button class="btn" type="submit">Загрузить</button>
        </div>
      </form>
    </div>
    <div class="card">
      <div style="font-weight:700; margin-bottom:8px;">Последние файлы</div>
      <ul>{items}</ul>
    </div>
  </div>
</body>
</html>"""

@app.route('/share/<doc_id>', methods=['GET', 'OPTIONS'])
def view_shared_doc(doc_id: str):
    if request.method == "OPTIONS":
        return ("", 204)
    ensure_upload_dir()
    target = None
    for name in os.listdir(UPLOAD_DIR):
        if name.startswith(f"{doc_id}__") and name.lower().endswith(".docx"):
            target = name
            break
    if not target:
        return "Файл не найден", 404

    path = os.path.join(UPLOAD_DIR, target)
    orig = target.split("__", 1)[1] if "__" in target else target
    html = docx_to_simple_html(path)
    download_url = f"/static/shared_docs/{target}"

    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{orig}</title>
  <style>
    body {{ font-family: Arial, sans-serif; background:#f0f2f5; margin:0; color:#111; }}
    .bar {{ background:#2c3e50; color:#fff; padding:14px 18px; display:flex; justify-content:space-between; gap:12px; flex-wrap:wrap; }}
    .wrap {{ max-width:1100px; margin:0 auto; padding:18px; }}
    .card {{ background:#fff; border-radius:12px; padding:16px; box-shadow:0 4px 6px rgba(0,0,0,0.05); }}
    a {{ color:#fff; text-decoration:none; opacity:0.9; }}
    .doc a {{ color:#2c3e50; text-decoration:underline; }}
    table {{ border-collapse:collapse; width:100%; margin:10px 0; }}
    td, th {{ border:1px solid #e6eaee; padding:8px; vertical-align:top; }}
  </style>
</head>
<body>
  <div class="bar">
    <div>{orig}</div>
    <div style="display:flex; gap:12px; align-items:center; flex-wrap:wrap;">
      <a href="/share">Назад</a>
      <a href="{download_url}">Скачать DOCX</a>
    </div>
  </div>
  <div class="wrap">
    <div class="card doc">{html}</div>
  </div>
</body>
</html>"""

@app.route('/api/save_tp_html', methods=['POST', 'OPTIONS'])
def save_tp_html():
    if request.method == "OPTIONS":
        return ("", 204)

    ensure_snapshot_dir()

    if 'file' not in request.files:
        return jsonify({"ok": False, "error": "file is required"}), 400
    f = request.files['file']
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "empty file"}), 400

    filename = secure_filename(f.filename)
    if not filename.lower().endswith(".html"):
        filename = filename + ".html"

    doc_id = f"tp-{int(__import__('time').time())}-{os.urandom(3).hex()}"
    save_name = f"{doc_id}__{filename}"
    path = os.path.join(SNAPSHOT_DIR, save_name)
    try:
        f.save(path)
    except Exception as e:
        return jsonify({"ok": False, "error": f"failed to save: {e}"}), 500

    url = f"/static/shared_docs/snapshots/{save_name}"
    return jsonify({"ok": True, "url": url, "id": doc_id, "name": filename})

@app.route('/api/upload_attachment', methods=['POST', 'OPTIONS'])
def upload_attachment():
    if request.method == "OPTIONS":
        return ("", 204)

    if 'file' not in request.files:
        return jsonify({"ok": False, "error": "file is required"}), 400
    f = request.files['file']
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "empty file"}), 400

    raw_filename = str(f.filename or "")
    raw_ext = os.path.splitext(raw_filename)[1].lower().lstrip(".")

    filename = secure_filename(raw_filename)
    ext = os.path.splitext(filename)[1].lower().lstrip(".")

    mimetype = str(getattr(f, "mimetype", "") or "").lower().strip()
    mime_to_ext = {
        "image/jpeg": "jpg",
        "image/jpg": "jpg",
        "image/png": "png",
        "image/gif": "gif",
        "image/webp": "webp",
        "image/bmp": "bmp",
        "image/tiff": "tiff",
        "image/svg+xml": "svg",
        "image/heic": "heic",
        "image/heif": "heif",
        "image/avif": "avif",
    }

    if not ext:
        ext = raw_ext or mime_to_ext.get(mimetype, "")
        if ext:
            filename = (filename or "file") + "." + ext
        elif not filename:
            filename = "file.bin"
            ext = "bin"

    if not filename:
        filename = f"file.{ext or 'bin'}"
        ext = ext or "bin"

    blocked = {
        "html", "htm", "js", "mjs", "css",
        "php", "py", "ps1", "bat", "cmd", "vbs",
        "exe", "dll", "com", "scr", "msi", "jar",
    }
    if ext in blocked and not mimetype.startswith("image/"):
        return jsonify({"ok": False, "error": f"unsupported file type: .{ext}"}), 400

    subdir = time.strftime("%Y%m%d", time.localtime())
    ensure_attachments_dir(subdir)


    file_id = f"att-{int(time.time())}-{os.urandom(3).hex()}"
    save_name = f"{file_id}__{filename}"
    path = os.path.join(ATTACHMENTS_DIR, subdir, save_name)
    try:
        f.save(path)
    except Exception as e:
        return jsonify({"ok": False, "error": f"failed to save: {e}"}), 500

    try:
        size = os.path.getsize(path)
    except Exception:
        size = None

    url = f"/static/attachments/{subdir}/{save_name}"
    return jsonify({"ok": True, "url": url, "name": filename, "ext": ext, "size": size})

@app.route('/api/parse_bom_docx', methods=['POST'])
def parse_bom_docx():
    #region debug-point auto-spec-recognition-backend
    def __dbg(event: str, payload: dict):
        try:
            import json
            import urllib.request

            url = "http://127.0.0.1:7778/event"
            body = json.dumps({
                "sessionId": "auto-spec-recognition",
                "event": str(event or "event"),
                "runId": "pre",
                "ts": int(time.time() * 1000),
                "payload": payload or {},
            }).encode("utf-8")
            req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"}, method="POST")
            urllib.request.urlopen(req, timeout=1).read()
        except Exception:
            pass
    #endregion

    if 'file' not in request.files:
        return jsonify({"ok": False, "error": "file is required"}), 400

    f = request.files['file']
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "empty file"}), 400

    filename = (f.filename or "").lower()
    if os.path.basename(filename).startswith("~$"):
        return jsonify({"ok": False, "error": "Это временный файл MS Word (~$...). Закройте документ в Word и выберите исходный .doc файл."}), 400
    if not (filename.endswith('.docx') or filename.endswith('.doc')):
        return jsonify({"ok": False, "error": "only .doc/.docx supported"}), 400

    __dbg("parse.enter", {"filename": str(f.filename or ""), "ext": (".docx" if filename.endswith(".docx") else ".doc")})

    try:
        from docx import Document
    except Exception:
        __dbg("parse.fail", {"stage": "import_docx", "error": "python-docx is not installed"})
        return jsonify({"ok": False, "error": "python-docx is not installed"}), 500

    if filename.endswith(".docx"):
        try:
            doc = Document(f.stream)
        except Exception as e:
            __dbg("parse.fail", {"stage": "read_docx", "error": str(e)})
            return jsonify({"ok": False, "error": f"failed to read docx: {e}"}), 400
    else:
        import shutil
        import subprocess
        import tempfile

        tmp_dir = tempfile.mkdtemp(prefix="tp_ctor_doc_")
        try:
            in_name = secure_filename(os.path.basename(f.filename)) or "spec.doc"
            if not in_name.lower().endswith(".doc"):
                in_name = f"{in_name}.doc"
            in_path = os.path.join(tmp_dir, in_name)
            try:
                f.save(in_path)
            except Exception as e:
                __dbg("convert.fail", {"stage": "save_doc", "error": str(e)})
                return jsonify({"ok": False, "error": f"failed to save .doc: {e}"}), 500

            out_path = os.path.splitext(in_path)[0] + ".docx"

            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice:
                for p in (
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                ):
                    if os.path.exists(p):
                        soffice = p
                        break
            __dbg("convert.detect", {"soffice": str(soffice or ""), "in_path": str(in_name)})
            if soffice:
                try:
                    subprocess.run(
                        [soffice, "--headless", "--nologo", "--convert-to", "docx", "--outdir", tmp_dir, in_path],
                        check=True,
                        capture_output=True,
                        text=True,
                        timeout=60,
                    )
                except subprocess.TimeoutExpired:
                    __dbg("convert.fail", {"stage": "soffice_timeout"})
                    return jsonify({"ok": False, "error": "DOC conversion timeout (soffice)"}), 500
                except Exception as e:
                    __dbg("convert.fail", {"stage": "soffice", "error": str(e)})
                    return jsonify({"ok": False, "error": f"DOC conversion failed (soffice): {e}"}), 500
            else:
                try:
                    import win32com.client  # type: ignore
                except Exception:
                    try:
                        import textwrap
                        vbs_path = os.path.join(tmp_dir, "convert_doc_to_docx.vbs")
                        vbs = textwrap.dedent(
                            r"""
                            On Error Resume Next
                            Dim inPath, outPath
                            inPath = WScript.Arguments(0)
                            outPath = WScript.Arguments(1)

                            Dim wordApp, doc
                            Set wordApp = CreateObject("Word.Application")
                            wordApp.Visible = False

                            Set doc = wordApp.Documents.Open(inPath, False, True)
                            If Err.Number <> 0 Then
                              WScript.Echo "open_failed:" & Err.Number & ":" & Err.Description
                              wordApp.Quit
                              WScript.Quit 2
                            End If

                            Err.Clear
                            doc.SaveAs2 outPath, 16
                            If Err.Number <> 0 Then
                              Err.Clear
                              doc.SaveAs outPath, 16
                            End If

                            If Err.Number <> 0 Then
                              WScript.Echo "save_failed:" & Err.Number & ":" & Err.Description
                              doc.Close False
                              wordApp.Quit
                              WScript.Quit 3
                            End If

                            doc.Close False
                            wordApp.Quit
                            WScript.Echo "ok"
                            WScript.Quit 0
                            """
                        ).strip()
                        with open(vbs_path, "w", encoding="utf-8") as fp:
                            fp.write(vbs)

                        r = subprocess.run(
                            ["cscript", "//nologo", vbs_path, in_path, out_path],
                            capture_output=True,
                            text=True,
                            timeout=60,
                        )
                        __dbg("convert.detect", {"method": "word_vbs", "rc": int(r.returncode), "stdout": (r.stdout or "")[:300], "stderr": (r.stderr or "")[:300]})
                        if r.returncode != 0:
                            __dbg("convert.fail", {"stage": "word_vbs", "rc": int(r.returncode), "stdout": (r.stdout or "")[:600], "stderr": (r.stderr or "")[:600]})
                            return jsonify({"ok": False, "error": "Не удалось конвертировать DOC через MS Word. Установите LibreOffice или используйте DOCX."}), 400
                    except Exception as e2:
                        __dbg("convert.fail", {"stage": "no_converter", "error": str(e2)})
                        return jsonify({"ok": False, "error": "Нужен конвертер .doc→.docx (LibreOffice/soffice или MS Word)"}), 400
                else:
                    word = None
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc_in = word.Documents.Open(in_path, ReadOnly=True)
                        try:
                            doc_in.SaveAs(out_path, FileFormat=16)
                        finally:
                            doc_in.Close(False)
                        word.Quit()
                        word = None
                    except Exception as e:
                        try:
                            if word:
                                word.Quit()
                        except Exception:
                            pass
                        __dbg("convert.fail", {"stage": "word", "error": str(e)})
                        return jsonify({"ok": False, "error": f"DOC conversion failed (Word): {e}"}), 500

            if not os.path.exists(out_path):
                found = [x for x in os.listdir(tmp_dir) if x.lower().endswith(".docx")]
                if found:
                    out_path = os.path.join(tmp_dir, found[0])
                else:
                    __dbg("convert.fail", {"stage": "no_docx_output"})
                    return jsonify({"ok": False, "error": "DOC conversion produced no .docx"}), 500

            try:
                doc = Document(out_path)
            except Exception as e:
                __dbg("convert.fail", {"stage": "read_converted", "error": str(e)})
                return jsonify({"ok": False, "error": f"failed to read converted docx: {e}"}), 400
        finally:
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass

    synonyms = {
        "level": ["уровень", "lvl", "level"],
        "pos": ["позиция", "поз.", "поз", "pos"],
        "designation": ["обозначение", "обознач.", "обозн.", "обозн", "обоз", "обозначние"],
        "code": ["код", "артикул", "марка", "designation", "code"],
        "name": ["наименование", "наим", "номенклатура", "деталь", "name"],
        "qty": ["кол-во", "количество", "кол.", "кол", "к-во", "qty", "count"],
        "unit": ["ед", "ед.", "ед.изм", "ед. изм", "единица", "unit"],
    }

    def norm(s: str) -> str:
        return " ".join((s or "").lower().replace("\xa0", " ").split()).strip()

    def find_col_index(headers, keys):
        for i, h in enumerate(headers):
            for k in keys:
                if k in h:
                    return i
        return -1

    def is_header_like(v: str) -> bool:
        n = norm(v)
        if not n:
            return False
        if n in ("поз", "поз.", "позиция", "обозначение", "код", "артикул", "наименование", "кол-во", "количество", "ед", "ед."):
            return True
        if n.startswith(("обознач", "обозн", "наимен", "кол")):
            return True
        return False

    def cell_indent_pt(cell) -> float:
        try:
            for p in getattr(cell, "paragraphs", []):
                if not (p.text or "").strip():
                    continue
                ind = getattr(getattr(p, "paragraph_format", None), "left_indent", None)
                if ind is None:
                    continue
                pt = getattr(ind, "pt", None)
                if pt is None:
                    continue
                return float(pt)
        except Exception:
            return 0.0
        return 0.0

    def infer_levels_from_indents(rows, code_idx: int, name_idx: int):
        indents = []
        for r in rows:
            ind = max(r.get("_indent_code", 0.0), r.get("_indent_name", 0.0))
            if ind > 0:
                indents.append(round(ind * 2) / 2.0)
        uniq = sorted(set(indents))
        if len(uniq) <= 1:
            return False

        for r in rows:
            ind = max(r.get("_indent_code", 0.0), r.get("_indent_name", 0.0))
            ind = round(ind * 2) / 2.0
            nearest = min(uniq, key=lambda x: abs(x - ind))
            r["level"] = uniq.index(nearest)
        return True

    def infer_levels_from_designation(rows):
        def depth(code: str) -> int:
            c = (code or "").strip()
            if not c:
                return 0
            if c.count(".") >= 1:
                return len([x for x in c.split(".") if x != ""])
            return 0

        depths = [depth(r.get("code")) for r in rows]
        non_zero = [d for d in depths if d > 0]
        if not non_zero:
            for r in rows:
                r["level"] = 0
            return

        base = min(non_zero)
        for r, d in zip(rows, depths):
            r["level"] = max(0, d - base) if d > 0 else 0

    def extract_product_name(document):
        import re

        lines = []
        for p in getattr(document, "paragraphs", []):
            raw = (p.text or "").replace("\xa0", " ").strip()
            if not raw:
                continue
            n = norm(raw)
            if not n:
                continue
            lines.append((raw, n))

        if not lines:
            return ""

        def after_sep(raw_line: str):
            for sep in (":", "—", "-"):
                if sep in raw_line:
                    right = raw_line.split(sep, 1)[1].strip()
                    if right:
                        return right
            return ""

        for raw, n in lines[:20]:
            if "наимен" in n or "издели" in n:
                val = after_sep(raw)
                if val and len(val) >= 3:
                    return val

        for raw, n in lines:
            if "спецификац" in n:
                cleaned = raw
                lower = n
                idx = lower.find("спецификац")
                if idx >= 0:
                    cleaned = raw[idx:]
                cleaned = cleaned.replace("СПЕЦИФИКАЦИЯ", "").replace("Спецификация", "").replace("спецификация", "").strip(" :—-")
                if cleaned:
                    return cleaned

        meta_bad = (
            "формат",
            "масштаб",
            "лист",
            "листов",
            "страниц",
            "инв",
            "подп",
            "утв",
            "разраб",
            "пров",
            "н.контр",
            "контр",
            "гост",
            "ту ",
            "ту\t",
            "стад",
            "лит",
            "материал",
            "масса",
        )

        code_re = re.compile(r"^[A-Za-zА-Яа-я0-9]{2,}(?:[./\\-][A-Za-zА-Яа-я0-9]{1,}){1,}[A-Za-zА-Яа-я0-9./\\-]{0,}$")
        best = ""
        best_score = -10
        for raw, n in lines[:30]:
            if any(x in n for x in meta_bad):
                continue
            if code_re.match(raw.strip()):
                continue
            if len(raw) < 4:
                continue
            if len(raw) > 140:
                continue
            score = 0
            if re.search(r"[А-Яа-я]", raw):
                score += 3
            if 6 <= len(raw) <= 90:
                score += 2
            if re.search(r"\d", raw):
                score += 1
            if score > best_score:
                best_score = score
                best = raw
        return best or lines[0][0]

    def extract_product_designation(document):
        import re

        lines = []
        for p in getattr(document, "paragraphs", []):
            raw = (p.text or "").replace("\xa0", " ").strip()
            if not raw:
                continue
            n = norm(raw)
            if not n:
                continue
            lines.append((raw, n))

        if not lines:
            return ""

        def after_sep(raw_line: str):
            for sep in (":", "—", "-"):
                if sep in raw_line:
                    right = raw_line.split(sep, 1)[1].strip()
                    if right:
                        return right
            return ""

        for raw, n in lines:
            if "обознач" in n or "шифр" in n or "код издел" in n or "обозн." in n:
                val = after_sep(raw)
                if val:
                    return val

        code_re = re.compile(r"([A-Za-zА-Яа-я0-9]{2,}(?:[./\\-][A-Za-zА-Яа-я0-9]{1,}){1,}[A-Za-zА-Яа-я0-9./\\-]{0,})")
        best = ""
        for raw, n in lines[:12]:
            if "спецификац" in n:
                continue
            for m in code_re.findall(raw):
                cand = (m or "").strip().strip(",.;")
                if len(cand) >= 4 and len(cand) > len(best):
                    best = cand
        return best

    candidates = []
    best_score = -1
    best_map = None
    best_prefix_text = ""
    best_header_score = -1

    for table_index, table in enumerate(doc.tables):
        matrix_norm = []
        matrix_raw = []
        matrix_indent = []
        for row in table.rows:
            row_norm = []
            row_raw = []
            row_indent = []
            for cell in row.cells:
                raw_text = (cell.text or "").replace("\xa0", " ")
                row_raw.append(raw_text.strip())
                row_norm.append(norm(raw_text))
                row_indent.append(cell_indent_pt(cell))
            matrix_norm.append(row_norm)
            matrix_raw.append(row_raw)
            matrix_indent.append(row_indent)

        if not matrix_norm:
            continue

        header_row_index = 0
        header_map = None
        score = -1

        for candidate in range(min(3, len(matrix_norm))):
            headers = matrix_norm[candidate]
            designation_idx = find_col_index(headers, synonyms["designation"])
            pos_idx = find_col_index(headers, synonyms["pos"])
            code_idx = designation_idx if designation_idx >= 0 else find_col_index(headers, synonyms["code"])
            if code_idx < 0 and pos_idx >= 0:
                code_idx = pos_idx

            m = {
                "level": find_col_index(headers, synonyms["level"]),
                "pos": pos_idx,
                "designation": designation_idx,
                "code": code_idx,
                "name": find_col_index(headers, synonyms["name"]),
                "qty": find_col_index(headers, synonyms["qty"]),
                "unit": find_col_index(headers, synonyms["unit"]),
            }

            found = sum(1 for v in (m["level"], m["code"], m["name"], m["qty"], m["unit"]) if v is not None and v >= 0)
            if m["designation"] is not None and m["designation"] >= 0:
                found += 3
            elif m["pos"] is not None and m["pos"] >= 0:
                found += 1
            if m["name"] >= 0:
                found += 2

            if found > score:
                score = found
                header_row_index = candidate
                header_map = m

        if not header_map:
            continue

        parsed_rows = []
        for r_raw, r_indent in zip(matrix_raw[header_row_index + 1:], matrix_indent[header_row_index + 1:]):
            def get(idx):
                return r_raw[idx] if idx is not None and idx >= 0 and idx < len(r_raw) else ""

            def get_indent(idx):
                return float(r_indent[idx]) if idx is not None and idx >= 0 and idx < len(r_indent) else 0.0

            code = get(header_map["code"]).strip()
            name = get(header_map["name"]).strip()
            qty = get(header_map["qty"]).strip()
            unit = get(header_map["unit"]).strip()
            level_raw = get(header_map["level"]).strip()

            if not code and not name:
                continue
            if is_header_like(code) and is_header_like(name):
                continue

            level = 0
            if header_map["level"] is not None and header_map["level"] >= 0 and level_raw and level_raw.isdigit():
                level = int(level_raw)

            qty_num = None
            if qty:
                try:
                    qty_num = float(qty.replace(",", "."))
                except Exception:
                    qty_num = None

            parsed_rows.append({
                "level": level,
                "code": code,
                "name": name,
                "qty": qty_num if qty_num is not None else qty,
                "unit": unit,
                "_indent_code": get_indent(header_map["code"]),
                "_indent_name": get_indent(header_map["name"]),
            })

        final_score = score * 10 + len(parsed_rows)
        if parsed_rows:
            prefix_cells = []
            for rr in matrix_raw[:header_row_index]:
                for c in rr:
                    cc = (c or "").strip()
                    if cc:
                        prefix_cells.append(cc)
            prefix_text = " ".join(prefix_cells).strip()
            candidates.append({
                "table_index": table_index,
                "rows": parsed_rows,
                "header_map": header_map,
                "header_score": score,
                "final_score": final_score,
                "prefix_text": prefix_text,
            })
            if final_score > best_score:
                best_score = final_score
                best_map = header_map
                best_prefix_text = prefix_text
                best_header_score = score

    if not candidates:
        __dbg("parse.fail", {"stage": "spec_table_not_found", "tables": len(getattr(doc, "tables", []) or [])})
        return jsonify({"ok": False, "error": "spec table not found"}), 404

    candidates.sort(key=lambda x: x["table_index"])
    selected = [c for c in candidates if c.get("header_score", -1) >= best_header_score - 1]
    if not selected:
        selected = candidates[:]

    all_rows = []
    for c in selected:
        all_rows.extend(c.get("rows") or [])

    if best_map and (best_map.get("level", -1) is None or best_map.get("level", -1) < 0):
        inferred = infer_levels_from_indents(all_rows, best_map.get("code", -1), best_map.get("name", -1))
        if not inferred:
            infer_levels_from_designation(all_rows)

    for r in all_rows:
        r.pop("_indent_code", None)
        r.pop("_indent_name", None)

    product_name = extract_product_name(doc)
    product_designation = extract_product_designation(doc)

    def looks_like_designation(v: str) -> bool:
        s = (v or "").strip()
        if not s:
            return False
        if s.lower() in ("поз", "поз.", "позиция", "обозначение", "обозн", "код", "наименование"):
            return False
        if len(s) < 3:
            return False
        if any(ch.isalnum() for ch in s) and any(sep in s for sep in (".", "-", "/", "\\")):
            return True
        if s.lower().endswith("сб") or s.lower().endswith("sb"):
            return True
        return False

    def is_generic_item_name(v: str) -> bool:
        n = norm(v)
        if not n:
            return True
        if "формат" in n:
            return True
        if any(x in n for x in ("чертеж", "схема", "перечень", "ведомост", "инструкц", "пояснит", "таблиц", "спецификац")):
            return True
        if n in ("упаковка",):
            return True
        if n in ("документация", "детали", "сборочные единицы", "стандартные изделия", "прочие изделия", "материалы", "комплекты"):
            return True
        return False

    best_code_row = None
    for r in all_rows:
        c = str(r.get("code") or "").strip()
        nm = str(r.get("name") or "").strip()
        if not c or not nm:
            continue
        if is_header_like(c) and is_header_like(nm):
            continue
        if not looks_like_designation(c):
            continue
        if is_generic_item_name(nm):
            continue
        best_code_row = r
        if c.lower().endswith("сб") or c.lower().endswith("sb"):
            break

    if best_code_row:
        if (not product_designation) or (product_designation.lower() in ("поз", "поз.", "позиция")):
            product_designation = str(best_code_row.get("code") or "").strip()
        if (not product_name) or is_generic_item_name(product_name):
            product_name = str(best_code_row.get("name") or "").strip()

    import re
    base = os.path.splitext(os.path.basename(str(f.filename or "")))[0]
    base = base.replace("\xa0", " ").strip()

    file_hint_code = ""
    file_hint_name = ""
    if "_" in base:
        parts = [p.strip() for p in base.split("_") if p.strip()]
        if parts:
            file_hint_code = parts[0]
        if len(parts) >= 2:
            file_hint_name = parts[1].split()[0].strip()

    if not file_hint_name:
        m = re.search(r"([A-Za-zА-Яа-я]{2,}-\d{1,6}(?:-\d{1,6})*)", base)
        if m:
            file_hint_name = m.group(1).strip()
        else:
            m2 = re.search(r"итс-\d{1,6}(?:-\d{1,6})*", norm(base))
            if m2:
                file_hint_name = m2.group(0).upper()

    if file_hint_code and looks_like_designation(file_hint_code):
        if (not product_designation) or (product_designation.lower() in ("поз", "поз.", "позиция")) or (len(file_hint_code) >= len(product_designation or "")):
            product_designation = file_hint_code

    if file_hint_name:
        if (not product_name) or is_generic_item_name(product_name) or (len(product_name) <= len(file_hint_name) + 2):
            product_name = file_hint_name

    if not product_designation and best_prefix_text:
        import re

        raw = best_prefix_text
        n = norm(raw)

        def after_sep(raw_line: str):
            for sep in (":", "—", "-"):
                if sep in raw_line:
                    right = raw_line.split(sep, 1)[1].strip()
                    if right:
                        return right
            return ""

        if "обознач" in n or "шифр" in n or "код издел" in n or "обозн." in n:
            val = after_sep(raw)
            if val:
                product_designation = val
        if not product_designation:
            code_re = re.compile(r"([A-Za-zА-Яа-я0-9]{2,}(?:[./\\-][A-Za-zА-Яа-я0-9]{1,}){1,}[A-Za-zА-Яа-я0-9./\\-]{0,})")
            best = ""
            for m in code_re.findall(raw):
                cand = (m or "").strip().strip(",.;")
                if len(cand) >= 4 and len(cand) > len(best):
                    best = cand
            product_designation = best

    if not product_name and best_prefix_text:
        raw = best_prefix_text
        cleaned = raw
        cleaned_norm = norm(raw)
        if "спецификац" in cleaned_norm:
            idx = cleaned_norm.find("спецификац")
            if idx >= 0:
                cleaned = raw[idx:]
        cleaned = cleaned.replace("СПЕЦИФИКАЦИЯ", "").replace("Спецификация", "").replace("спецификация", "").strip(" :—-")
        if cleaned:
            product_name = cleaned

    if not product_designation and all_rows:
        product_designation = (all_rows[0].get("code") or "").strip()
    if not product_name and all_rows:
        product_name = (all_rows[0].get("name") or "").strip()
    __dbg("parse.ok", {"rows": len(all_rows), "tables_used": len(selected), "product_name": str(product_name or ""), "product_designation": str(product_designation or "")})
    return jsonify({"ok": True, "rows": all_rows, "columns": best_map, "product_name": product_name, "product_designation": product_designation, "tables_used": len(selected)})

@app.route('/api/parse_bom_excel', methods=['POST'])
def parse_bom_excel():
    if 'file' not in request.files:
        return jsonify({"ok": False, "error": "file is required"}), 400

    f = request.files['file']
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "empty file"}), 400

    filename = (f.filename or "").lower()
    if not (filename.endswith('.xlsx') or filename.endswith('.xlsm')):
        return jsonify({"ok": False, "error": "only .xlsx/.xlsm supported"}), 400

    try:
        import openpyxl
    except Exception:
        return jsonify({"ok": False, "error": "openpyxl is not installed"}), 500

    try:
        wb = openpyxl.load_workbook(f.stream, data_only=True)
    except Exception as e:
        return jsonify({"ok": False, "error": f"failed to read excel: {e}"}), 400

    synonyms = {
        "level": ["уровень", "lvl", "level"],
        "pos": ["позиция", "поз.", "поз", "pos"],
        "designation": ["обозначение", "обознач.", "обозн.", "обозн", "обоз", "обозначние"],
        "code": ["код", "артикул", "марка", "designation", "code"],
        "name": ["наименование", "наим", "номенклатура", "деталь", "name"],
        "qty": ["кол-во", "количество", "кол.", "кол", "к-во", "qty", "count"],
        "unit": ["ед", "ед.", "ед.изм", "ед. изм", "единица", "unit"],
    }

    def norm(s: str) -> str:
        return " ".join((s or "").lower().replace("\xa0", " ").split()).strip()

    def as_text(v) -> str:
        if v is None:
            return ""
        if isinstance(v, str):
            return v.strip()
        try:
            if isinstance(v, (int, float)):
                if isinstance(v, bool):
                    return str(v)
                if abs(v - int(v)) < 1e-9:
                    return str(int(v))
                return str(v)
        except Exception:
            pass
        return str(v).strip()

    def find_col_index(headers, keys):
        for i, h in enumerate(headers):
            for k in keys:
                if k in h:
                    return i
        return -1

    def find_code_col(headers_norm):
        primary = ["обозначение", "обознач.", "обозн", "обозн.", "обоз", "обозначние", "код", "артикул", "марка", "designation", "part", "code"]
        secondary = ["позиция", "поз.", "поз", "pos"]
        idx = find_col_index(headers_norm, primary)
        if idx >= 0:
            return idx
        return find_col_index(headers_norm, secondary)

    def is_header_like(v: str) -> bool:
        n = norm(v)
        if not n:
            return False
        if n in ("поз", "поз.", "позиция", "обозначение", "код", "артикул", "наименование", "кол-во", "количество", "ед", "ед."):
            return True
        if n.startswith(("обознач", "обозн", "наимен", "кол")):
            return True
        return False

    max_scan = 80

    candidates = []
    best_sheet_score = -1

    for ws in list(getattr(wb, "worksheets", []) or []):
        rows_raw = []
        for r in ws.iter_rows(min_row=1, max_row=max_scan, values_only=True):
            rows_raw.append([as_text(v) for v in (r or [])])

        sheet_best = None
        sheet_best_score = -1
        sheet_best_map = None

        for idx in range(min(20, len(rows_raw))):
            headers_norm = [norm(x) for x in rows_raw[idx]]
            m = {
                "level": find_col_index(headers_norm, synonyms["level"]),
                "code": find_code_col(headers_norm),
                "name": find_col_index(headers_norm, synonyms["name"]),
                "qty": find_col_index(headers_norm, synonyms["qty"]),
                "unit": find_col_index(headers_norm, synonyms["unit"]),
            }

            found = sum(1 for v in m.values() if v >= 0)
            if m["code"] >= 0:
                found += 2
            if m["name"] >= 0:
                found += 2
            if m["qty"] >= 0:
                found += 1
            if m["unit"] >= 0:
                found += 1

            if found > sheet_best_score:
                sheet_best_score = found
                sheet_best = idx
                sheet_best_map = m

        if sheet_best is None or sheet_best_map is None or sheet_best_map.get("code", -1) < 0 or sheet_best_map.get("name", -1) < 0:
            continue

        code_idx = sheet_best_map.get("code", -1)
        name_idx = sheet_best_map.get("name", -1)

        non_empty = 0
        non_empty_code = 0
        for r in ws.iter_rows(min_row=sheet_best + 2, max_row=sheet_best + 42, values_only=True):
            cells = [as_text(v) for v in (r or [])]
            code = cells[code_idx].strip() if code_idx >= 0 and code_idx < len(cells) else ""
            name = cells[name_idx].strip() if name_idx >= 0 and name_idx < len(cells) else ""
            if code or name:
                non_empty += 1
            if code:
                non_empty_code += 1

        total_score = float(sheet_best_score) + min(non_empty, 20) / 2.0 + min(non_empty_code, 20)

        candidates.append({
            "ws": ws,
            "rows_raw": rows_raw,
            "header_row_index": sheet_best,
            "map": sheet_best_map,
            "header_score": sheet_best_score,
            "score": total_score,
        })
        if total_score > best_sheet_score:
            best_sheet_score = total_score

    if not candidates:
        return jsonify({"ok": False, "error": "spec header not found"}), 404

    candidates.sort(key=lambda x: x["score"], reverse=True)
    lead = candidates[0]
    lead_score = float(lead.get("score") or 0.0)
    selected = [c for c in candidates if float(c.get("score") or 0.0) >= lead_score - 2.0]
    if not selected:
        selected = [lead]

    parsed_rows = []
    sheets_used = []

    for c in selected:
        ws = c["ws"]
        header_row_index = c["header_row_index"]
        m = c["map"] or {}
        code_idx = m.get("code", -1)
        name_idx = m.get("name", -1)
        qty_idx = m.get("qty", -1)
        unit_idx = m.get("unit", -1)
        level_idx = m.get("level", -1)

        if header_row_index is None or code_idx < 0 or name_idx < 0:
            continue

        local_rows = []
        for r in ws.iter_rows(min_row=header_row_index + 2, values_only=True):
            cells = [as_text(v) for v in (r or [])]
            code = cells[code_idx].strip() if code_idx >= 0 and code_idx < len(cells) else ""
            name = cells[name_idx].strip() if name_idx >= 0 and name_idx < len(cells) else ""
            if not code and not name:
                continue
            if is_header_like(code) and is_header_like(name):
                continue

            qty = cells[qty_idx].strip() if qty_idx >= 0 and qty_idx < len(cells) else ""
            unit = cells[unit_idx].strip() if unit_idx >= 0 and unit_idx < len(cells) else ""

            qty_num = None
            if qty:
                try:
                    qty_num = float(qty.replace(",", "."))
                except Exception:
                    qty_num = None

            level = ""
            if level_idx >= 0 and level_idx < len(cells):
                level = cells[level_idx].strip()

            local_rows.append({
                "level": level,
                "code": code,
                "name": name,
                "qty": qty_num if qty_num is not None else qty,
                "unit": unit,
            })

        if local_rows:
            parsed_rows.extend(local_rows)
            sheets_used.append(ws.title)

    def infer_levels_from_designation(rows):
        def depth(code: str) -> int:
            c = (code or "").strip()
            if not c:
                return 0
            if c.count(".") >= 1:
                return len([x for x in c.split(".") if x != ""])
            return 0

        depths = [depth(r.get("code")) for r in rows]
        non_zero = [d for d in depths if d > 0]
        if not non_zero:
            for r in rows:
                r["level"] = 0
            return

        base = min(non_zero)
        for r, d in zip(rows, depths):
            r["level"] = max(0, d - base) if d > 0 else 0

    def normalize_levels(rows):
        has_level = any(str(r.get("level") or "").strip() != "" for r in rows)
        if not has_level:
            infer_levels_from_designation(rows)
            return

        for r in rows:
            v = str(r.get("level") or "").strip()
            try:
                r["level"] = int(float(v)) if v != "" else 0
            except Exception:
                r["level"] = 0

    normalize_levels(parsed_rows)

    first_code_row = None
    for r in parsed_rows:
        c = str(r.get("code") or "").strip()
        if c:
            first_code_row = r
            if c.lower().endswith("сб") or c.lower().endswith("sb"):
                break

    product_name = ""
    product_designation = ""
    if first_code_row:
        product_designation = str(first_code_row.get("code") or "").strip()
        product_name = str(first_code_row.get("name") or "").strip()

    rows_raw = lead.get("rows_raw") or []
    header_row_index = lead.get("header_row_index") or 0
    for i in range(max(0, header_row_index - 12), header_row_index + 1):
        row = rows_raw[i] if i < len(rows_raw) else []
        joined = " ".join([x for x in row if x]).strip()
        n = norm(joined)
        if not product_designation and ("обознач" in n or "шифр" in n or "код издел" in n or "обозн." in n):
            for sep in (":", "—", "-"):
                if sep in joined:
                    right = joined.split(sep, 1)[1].strip()
                    if right and len(right) >= 4:
                        product_designation = right
                        break
        if not product_name and "спецификац" in n:
            cleaned = joined.replace("СПЕЦИФИКАЦИЯ", "").replace("Спецификация", "").replace("спецификация", "").strip(" :—-")
            if cleaned:
                product_name = cleaned

    if parsed_rows:
        first_with_code = None
        for r in parsed_rows:
            c = str(r.get("code") or "").strip()
            n = str(r.get("name") or "").strip()
            if c:
                first_with_code = (c, n)
                if c.lower().endswith("сб") or c.lower().endswith("sb"):
                    break

        if not product_designation:
            product_designation = first_with_code[0] if first_with_code else ""
        if not product_name:
            if first_with_code and first_with_code[1]:
                product_name = first_with_code[1]
            else:
                product_name = (parsed_rows[0].get("name") or "").strip()

    return jsonify({
        "ok": True,
        "rows": parsed_rows,
        "columns": lead.get("map"),
        "product_name": product_name,
        "product_designation": product_designation,
        "sheet": (sheets_used[0] if sheets_used else lead["ws"].title),
        "sheets": sheets_used
    })

@app.route('/api/export_bom_excel', methods=['POST'])
def export_bom_excel():
    try:
        import openpyxl
    except Exception:
        return jsonify({"ok": False, "error": "openpyxl is not installed"}), 500

    data = request.get_json(silent=True) or {}
    columns = data.get("columns")
    rows = data.get("rows")
    filename = secure_filename(str(data.get("filename") or "spec.xlsx")) or "spec.xlsx"
    if not filename.lower().endswith(".xlsx"):
        filename = f"{filename}.xlsx"

    if not isinstance(columns, list) or not columns:
        return jsonify({"ok": False, "error": "columns is required"}), 400
    if not isinstance(rows, list):
        return jsonify({"ok": False, "error": "rows is required"}), 400

    header = []
    for c in columns:
        if isinstance(c, dict):
            header.append(str(c.get("label") or c.get("key") or ""))
        else:
            header.append(str(c))
    if not any(h.strip() for h in header):
        header = ["Обозначение", "Наименование", "Кол-во", "Ед."]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Спецификация"
    ws.append(header)
    try:
        for r in rows:
            if isinstance(r, list):
                ws.append(r[: len(header)] + [""] * max(0, len(header) - len(r)))
            else:
                ws.append([str(r)])
    except Exception as e:
        return jsonify({"ok": False, "error": f"failed to build xlsx: {e}"}), 400

    try:
        from io import BytesIO
        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
    except Exception as e:
        return jsonify({"ok": False, "error": f"failed to save xlsx: {e}"}), 500

    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route('/api/export_bom_doc', methods=['POST'])
def export_bom_doc():
    data = request.get_json(silent=True) or {}
    columns = data.get("columns")
    rows = data.get("rows")
    filename = secure_filename(str(data.get("filename") or "spec.doc")) or "spec.doc"
    if not filename.lower().endswith(".doc"):
        filename = f"{filename}.doc"

    if not isinstance(columns, list) or not columns:
        return jsonify({"ok": False, "error": "columns is required"}), 400
    if not isinstance(rows, list):
        return jsonify({"ok": False, "error": "rows is required"}), 400

    norm_cols = []
    for c in columns:
        if isinstance(c, dict):
            norm_cols.append({"key": str(c.get("key") or ""), "label": str(c.get("label") or c.get("key") or "")})
        else:
            norm_cols.append({"key": str(c), "label": str(c)})

    key_to_idx = {c["key"]: i for i, c in enumerate(norm_cols) if c.get("key")}
    level_idx = key_to_idx.get("level", -1)
    code_idx = key_to_idx.get("code", -1)
    name_idx = key_to_idx.get("name", -1)
    qty_idx = key_to_idx.get("qty", -1)
    unit_idx = key_to_idx.get("unit", -1)

    def rtf_escape(s: str) -> str:
        s = (s or "")
        s = s.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
        s = s.replace("\r\n", "\n").replace("\r", "\n")
        s = s.replace("\n", "\\line ")
        return s

    header = ["Обозначение", "Наименование", "Кол-во", "Ед."]
    if code_idx < 0 or name_idx < 0:
        header = [c["label"] for c in norm_cols if c.get("label")]

    lines = []
    lines.append(rtf_escape("\t".join(header)) + "\\line ")

    for r in rows:
        if not isinstance(r, list):
            continue
        level = 0
        if level_idx >= 0 and level_idx < len(r):
            try:
                v = str(r[level_idx] or "").strip()
                level = int(float(v)) if v != "" else 0
            except Exception:
                level = 0
        code = str(r[code_idx] if code_idx >= 0 and code_idx < len(r) else "").strip()
        name = str(r[name_idx] if name_idx >= 0 and name_idx < len(r) else "").strip()
        qty = str(r[qty_idx] if qty_idx >= 0 and qty_idx < len(r) else "").strip()
        unit = str(r[unit_idx] if unit_idx >= 0 and unit_idx < len(r) else "").strip()
        if level > 0 and name:
            name = ("    " * min(level, 12)) + name
        lines.append(rtf_escape("\t".join([code, name, qty, unit])) + "\\line ")

    rtf = (
        "{\\rtf1\\ansi\\ansicpg1251\\deff0"
        "{\\fonttbl{\\f0 Calibri;}}"
        "\\fs22 "
        + "".join(lines)
        + "}"
    )

    from io import BytesIO

    bio = BytesIO(rtf.encode("cp1251", errors="replace"))
    bio.seek(0)
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/msword",
    )

@app.route('/api/status')
def get_status():
    # Базовая заглушка для статуса автомата
    return jsonify({
        "status": "online",
        "message": "Система готова к работе",
        "process": "Ожидание команды"
    })

if __name__ == '__main__':
    host = os.getenv("HOST", "0.0.0.0")
    preferred_port = int(os.getenv("PORT", "5000"))
    candidate_ports = [preferred_port, 8000, 8080, 5000]

    last_error = None
    for port in candidate_ports:
        try:
            app.run(debug=True, host=host, port=port, use_reloader=False)
            break
        except OSError as e:
            last_error = e
            continue

    if last_error:
        raise last_error
